import os
import difflib
import re
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext


class ExamComparatorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("试卷题目查重工具 v2.0")
        self.root.geometry("900x650")

        self.create_widgets()
        self.files = []

    def create_widgets(self):
        # 文件选择区域
        file_frame = tk.LabelFrame(self.root, text="选择试卷文件 (支持多选)", padx=5, pady=5)
        file_frame.pack(fill="x", padx=10, pady=5)

        self.file_listbox = tk.Listbox(file_frame, height=6, selectmode=tk.EXTENDED)
        self.file_listbox.pack(fill="x", expand=True, pady=5)

        btn_frame = tk.Frame(file_frame)
        btn_frame.pack(fill="x", pady=5)

        tk.Button(btn_frame, text="添加文件", command=self.add_files).pack(side="left", padx=5)
        tk.Button(btn_frame, text="移除选中", command=self.remove_file).pack(side="left", padx=5)
        tk.Button(btn_frame, text="清空列表", command=self.clear_files).pack(side="left", padx=5)

        # 操作按钮
        action_frame = tk.Frame(self.root)
        action_frame.pack(fill="x", padx=10, pady=5)

        tk.Button(action_frame, text="开始查重", command=self.run_comparison, bg="#4CAF50", fg="white").pack(
            side="left", padx=5, ipadx=10)
        tk.Button(action_frame, text="保存报告", command=self.save_report, bg="#2196F3", fg="white").pack(side="left",
                                                                                                          padx=5,
                                                                                                          ipadx=10)
        tk.Button(action_frame, text="退出", command=self.root.quit, bg="#f44336", fg="white").pack(side="right",
                                                                                                    padx=5, ipadx=10)

        # 结果展示区域
        result_frame = tk.LabelFrame(self.root, text="查重结果", padx=5, pady=5)
        result_frame.pack(fill="both", expand=True, padx=10, pady=5)

        self.result_text = scrolledtext.ScrolledText(result_frame, wrap=tk.WORD, font=('微软雅黑', 10))
        self.result_text.pack(fill="both", expand=True)

    def add_files(self):
        files = filedialog.askopenfilenames(
            title="选择试卷文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if files:
            for f in files:
                if f not in self.files:
                    self.files.append(f)
                    self.file_listbox.insert(tk.END, f"{os.path.basename(f)} ({os.path.dirname(f)})")

    def remove_file(self):
        selection = self.file_listbox.curselection()
        if selection:
            for i in reversed(selection):  # 从后往前删除避免索引变化
                self.file_listbox.delete(i)
                del self.files[i]

    def clear_files(self):
        self.file_listbox.delete(0, tk.END)
        self.files = []

    def run_comparison(self):
        if len(self.files) < 2:
            messagebox.showerror("错误", "请至少选择2个文件进行比较")
            return

        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, "正在分析，请稍候...\n")
        self.root.update()

        try:
            results = self.compare_documents(self.files)
            report = self.create_report(results)
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, report)
            messagebox.showinfo("完成", "查重分析已完成!")
        except Exception as e:
            messagebox.showerror("错误", f"分析过程中出错: {str(e)}")

    def save_report(self):
        report = self.result_text.get(1.0, tk.END)
        if not report.strip():
            messagebox.showerror("错误", "没有可保存的内容")
            return

        file_path = filedialog.asksaveasfilename(
            title="保存查重报告",
            defaultextension=".txt",
            filetypes=[("文本文件", "*.txt"), ("Word文档", "*.docx"), ("所有文件", "*.*")]
        )

        if file_path:
            try:
                if file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_paragraph(report)
                    doc.save(file_path)
                else:
                    with open(file_path, "w", encoding="utf-8") as f:
                        f.write(report)
                messagebox.showinfo("成功", f"报告已保存到:\n{file_path}")
            except Exception as e:
                messagebox.showerror("错误", f"保存失败: {str(e)}")

    @staticmethod
    def remove_options(question_text):
        """改进的选择题选项去除方法"""
        # 处理A. B. C. D. 或 A． B． C． D．等格式
        pattern = r'^[A-E][\.．、]\s*.+$'
        lines = question_text.split('\n')
        cleaned_lines = []

        for line in lines:
            if not re.match(pattern, line.strip()):
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    @staticmethod
    def extract_questions(doc_path):
        """改进的题目提取方法"""
        doc = Document(doc_path)
        questions = []
        current_question = ""
        question_num = 0
        last_is_question = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 改进的题目编号检测，支持多种格式: 1. 1、 (1) 一、 等
            match = re.match(r'^(\d+)[\.．、]|^[(（](\d+)[)）]|^[一二三四五六七八九十]+[、．.]', text)
            if match:
                if current_question:  # 保存上一题
                    questions.append((question_num, ExamComparatorApp.remove_options(current_question)))

                # 确定题目编号
                if match.group(1):  # 数字编号 1. 或 1、
                    question_num = int(match.group(1))
                elif match.group(2):  # (1) 或 (1)
                    question_num = int(match.group(2))
                else:  # 中文编号 一、 二、等
                    chinese_num = match.group(0)[0]
                    num_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
                               '十': 10}
                    question_num = num_map.get(chinese_num, 0)

                current_question = text
                last_is_question = True
            else:
                # 如果上一行是题目，这一行可能是题目内容或选项
                if last_is_question:
                    current_question += "\n" + text
                else:
                    # 可能是没有编号的题目(如大题中的小题)
                    if current_question:
                        current_question += "\n" + text
                    else:
                        current_question = text
                last_is_question = False

        if current_question:
            questions.append((question_num, ExamComparatorApp.remove_options(current_question)))

        return questions

    @staticmethod
    def calculate_similarity(text1, text2):
        """改进的相似度计算方法"""
        # 去除多余空白字符
        text1 = re.sub(r'\s+', ' ', text1.strip())
        text2 = re.sub(r'\s+', ' ', text2.strip())
        return difflib.SequenceMatcher(None, text1, text2).ratio()

    @staticmethod
    def compare_documents(files):
        """比较多个文档的题目"""
        all_questions = {}
        for file in files:
            filename = os.path.basename(file)
            questions = ExamComparatorApp.extract_questions(file)
            all_questions[filename] = questions

        results = []

        # 两两比较文档
        filenames = list(all_questions.keys())
        for i in range(len(filenames)):
            for j in range(i + 1, len(filenames)):
                file1 = filenames[i]
                file2 = filenames[j]
                q1_list = all_questions[file1]
                q2_list = all_questions[file2]

                duplicates = []
                similar = []

                for num1, text1 in q1_list:
                    for num2, text2 in q2_list:
                        sim = ExamComparatorApp.calculate_similarity(text1, text2)
                        if sim == 1.0:
                            duplicates.append((num1, num2, text1))
                        elif sim > 0.9:
                            similar.append((num1, num2, text1, text2, sim))

                total_q1 = len(q1_list)
                total_q2 = len(q2_list)
                dup_count = len(duplicates)
                sim_count = len(similar)
                dup_rate = dup_count / min(total_q1, total_q2) if min(total_q1, total_q2) > 0 else 0
                sim_rate = (dup_count + sim_count) / min(total_q1, total_q2) if min(total_q1, total_q2) > 0 else 0

                results.append({
                    'file1': file1,
                    'file2': file2,
                    'duplicates': duplicates,
                    'similar': similar,
                    'dup_count': dup_count,
                    'sim_count': sim_count,
                    'dup_rate': dup_rate,
                    'sim_rate': sim_rate,
                    'total_q1': total_q1,
                    'total_q2': total_q2
                })

        return results

    @staticmethod
    def create_report(results):
        """改进的报告生成方法"""
        report = "=" * 80 + "\n"
        report += "试卷题目查重报告\n"
        report += "=" * 80 + "\n\n"

        # 汇总统计
        total_comparisons = len(results)
        total_dup = sum(r['dup_count'] for r in results)
        total_sim = sum(r['sim_count'] for r in results)

        report += f"【汇总统计】\n"
        report += f"比较试卷对数: {total_comparisons}\n"
        report += f"完全重复题目总数: {total_dup}\n"
        report += f"高度相似题目总数: {total_sim}\n\n"

        # 详细比较结果
        for i, result in enumerate(results, 1):
            report += f"【比较组 {i}】\n"
            report += f"对比文件1: {result['file1']} (共{result['total_q1']}题)\n"
            report += f"对比文件2: {result['file2']} (共{result['total_q2']}题)\n"
            report += "-" * 60 + "\n"
            report += f"完全重复题目数量: {result['dup_count']} (重复率: {result['dup_rate']:.2%})\n"
            report += f"高度相似题目数量: {result['sim_count']} (相似率: {result['sim_rate']:.2%})\n"

            if result['duplicates']:
                report += "\n◆ 完全重复题目:\n"
                for dup in result['duplicates']:
                    report += f"├─ {result['file1']} 第{dup[0]}题 ←→ {result['file2']} 第{dup[1]}题\n"
                    report += f"│  题目内容: {dup[2][:120]}{'...' if len(dup[2]) > 120 else ''}\n"

            if result['similar']:
                report += "\n◆ 高度相似题目(相似度>90%):\n"
                for sim in result['similar']:
                    report += f"├─ {result['file1']} 第{sim[0]}题 ←→ {result['file2']} 第{sim[1]}题 (相似度: {sim[4]:.2%})\n"
                    report += f"│  {result['file1']}题目: {sim[2][:60]}{'...' if len(sim[2]) > 60 else ''}\n"
                    report += f"│  {result['file2']}题目: {sim[3][:60]}{'...' if len(sim[3]) > 60 else ''}\n"

            report += "\n" + "=" * 80 + "\n\n"

        return report


if __name__ == "__main__":
    root = tk.Tk()
    app = ExamComparatorApp(root)
    root.mainloop()